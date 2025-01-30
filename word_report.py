import sqlite3
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os
from datetime import datetime
from dateutil.relativedelta import relativedelta
import calendar
import logging

class DatabaseConnection:
    def __init__(self, db_path):
        self.db_path = db_path
        self.conn = None
        
    def __enter__(self):
        self.conn = sqlite3.connect(self.db_path)
        return self.conn
        
    def __exit__(self, exc_type, exc_val, exc_tb):
        if self.conn:
            self.conn.close()

class WordReportGenerator:
    def __init__(self, db_path):
        logging.info(f"Initialiserer WordReportGenerator med database: {db_path}")
        self.db_path = db_path
        try:
            self.min_km = self.get_min_km_setting()
            logging.info(f"Minimum kilometer sat til: {self.min_km}")
        except Exception as e:
            logging.error(f"Fejl ved initialisering af WordReportGenerator: {str(e)}")
            raise
        self.doc = Document()
        
        # Definer kolonne grupper
        self.driftsdata_kolonner = [
            'Ø Forbrug [l/100km]',
            'Ø Rækkevidde ved forbrug [km/l]',
            'Ø Forbrug ved kørsel [l/100km]',
            'Forbrug [l]',
            'Kørestrækning [km]',
            'Ø totalvægt [t]'
        ]
        
        self.korselsdata_kolonner = [
            'Aktiv påløbsdrift (km) [km]',
            'Afstand i påløbsdrift [km]',
            'Kickdown (km) [km]',
            'Afstand med kørehastighedsregulering (> 50 km/h) [km]',
            'Afstand > 50 km/h uden kørehastighedsregulering [km]',
            'Forbrug med kørehastighedsregulering [l/100km]',
            'Forbrug uden kørehastighedsregulering [l/100km]',
            'Driftsbremse (km) [km]',
            'Afstand motorbremse [km]',
            'Overspeed (km uden påløbsdrift) [km]'
        ]
        
        self.tomgangsdata_kolonner = [
            'Motordriftstid [hh:mm:ss]',
            'Køretid [hh:mm:ss]',
            'Tomgang / stilstandstid [hh:mm:ss]'
        ]

    def get_min_km_setting(self):
        """Henter minimum kilometer indstilling fra databasen"""
        try:
            with DatabaseConnection('databases/settings.db') as conn:
                cursor = conn.cursor()
                cursor.execute('SELECT value FROM settings WHERE key = "min_km"')
                result = cursor.fetchone()
                return float(result[0]) if result else 100.0
        except:
            return 100.0

    def konverter_tid_til_sekunder(self, tid_str):
        """Konverterer tid i format 'tt:mm:ss' til sekunder"""
        try:
            t, m, s = tid_str.split(':')
            return int(t) * 3600 + int(m) * 60 + int(s)
        except:
            return 0

    def beregn_noegletal(self, data):
        """Beregner nøgletal baseret på kørselsdata"""
        try:
            noegletal = {}
            
            # Beregn tomgangsprocent
            motor_tid = self.konverter_tid_til_sekunder(data['Motordriftstid [hh:mm:ss]'])
            tomgangs_tid = self.konverter_tid_til_sekunder(data['Tomgang / stilstandstid [hh:mm:ss]'])
            noegletal['Tomgangsprocent'] = (tomgangs_tid / motor_tid) * 100 if motor_tid > 0 else 0
            
            # Beregn cruise control andel
            distance_over_50 = float(data['Afstand med kørehastighedsregulering (> 50 km/h) [km]']) + \
                             float(data['Afstand > 50 km/h uden kørehastighedsregulering [km]'])
            cruise_distance = float(data['Afstand med kørehastighedsregulering (> 50 km/h) [km]'])
            noegletal['Fartpilot Andel'] = (cruise_distance / distance_over_50) * 100 if distance_over_50 > 0 else 0
            
            # Beregn motorbremse andel
            driftsbremse = float(data['Driftsbremse (km) [km]'])
            motorbremse = float(data['Afstand motorbremse [km]'])
            total_bremse_distance = driftsbremse + motorbremse
            noegletal['Motorbremse Andel'] = (motorbremse / total_bremse_distance) * 100 if total_bremse_distance > 0 else 0
            
            # Beregn diesel forbrug pr. km/l
            forbrug_liter = float(data['Forbrug [l]'])
            korestraekning = float(data['Kørestrækning [km]'])
            if forbrug_liter > 0:
                noegletal['Diesel Effektivitet'] = korestraekning / forbrug_liter
            else:
                noegletal['Diesel Effektivitet'] = 0
            
            # Beregn vægtkorrigeret forbrug
            total_vaegt = float(data.get('Ø totalvægt [t]', 0))
            if total_vaegt > 0 and forbrug_liter > 0:
                vaegt_korrigeret_forbrug = (forbrug_liter / korestraekning * 100) / total_vaegt
                noegletal['Vægtkorrigeret Forbrug'] = vaegt_korrigeret_forbrug
            else:
                noegletal['Vægtkorrigeret Forbrug'] = 0

            # Beregn påløbsdrift andel
            total_distance = float(data['Kørestrækning [km]'])
            aktiv_paalobsdrift = float(data['Aktiv påløbsdrift (km) [km]'])
            afstand_paalobsdrift = float(data['Afstand i påløbsdrift [km]'])
            if total_distance > 0:
                noegletal['Påløbsdrift Andel'] = ((aktiv_paalobsdrift + afstand_paalobsdrift) / total_distance) * 100
            else:
                noegletal['Påløbsdrift Andel'] = 0
            
            # Beregn Overspeed andel
            overspeed = float(data['Overspeed (km uden påløbsdrift) [km]'])
            if total_distance > 0:
                noegletal['Overspeed Andel'] = (overspeed / total_distance) * 100
            else:
                noegletal['Overspeed Andel'] = 0
            
            return noegletal
            
        except Exception as e:
            print(f"Fejl ved beregning af nøgletal: {str(e)}")
            return {}

    def opret_samlet_rangering(self, kvalificerede_chauffoerer):
        """Opretter en samlet rangering baseret på de fire hovedparametre"""
        self.tilfoej_sektion_overskrift("Samlet Performance Rangering")
        
        intro_tekst = self.doc.add_paragraph()
        intro_run = intro_tekst.add_run(
            "Den samlede rangering kombinerer præstationen på fire nøgleområder med følgende virksomhedsmål:\n\n"
            "1. Tomgang: Mål på max 5% - Minimering af unødvendig tomgangskørsel\n"
            "2. Fartpilot: Mål på minimum 66,5% - Optimal brug af fartpilot ved højere hastigheder\n"
            "3. Motorbremse: Mål på minimum 56% - Effektiv brug af motorbremsning\n"
            "4. Påløbsdrift: Mål på minimum 7% - Udnyttelse af køretøjets momentum\n\n"
            "Hver chauffør får points baseret på deres placering i hver kategori. "
            "Lavere samlet score er bedre, da det betyder bedre placeringer på tværs af kategorierne. "
            "De tre bedste chauffører er markeret med grøn for at fremhæve særligt god præstation.\n"
            "Målene er sat af virksomheden og bruges som reference for optimal kørsel.\n"
        )
        intro_run.font.size = Pt(11)
        
        # Hent data for alle kvalificerede chauffører
        chauffoer_data = {}
        noegletal_data = {}
        
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        for chauffoer, _ in kvalificerede_chauffoerer:
            cursor.execute('SELECT * FROM chauffør_data_data WHERE Chauffør = ?', (chauffoer,))
            data = dict(zip([col[0] for col in cursor.description], cursor.fetchone()))
            chauffoer_data[chauffoer] = data
            noegletal_data[chauffoer] = self.beregn_noegletal(data)
        
        conn.close()
        
        # Definer nøgletal til rangering
        noegletal_ranking = {
            'Tomgangsprocent': False,
            'Fartpilot Andel': True,
            'Motorbremse Andel': True,
            'Påløbsdrift Andel': True
        }
        
        # Beregn placering for hver chauffør i hver kategori
        placeringer = {chauffoer: {} for chauffoer in noegletal_data.keys()}
        
        for noegletal, hoejere_er_bedre in noegletal_ranking.items():
            sorterede_chauffoerer = sorted(
                noegletal_data.items(),
                key=lambda x: x[1].get(noegletal, 0),
                reverse=hoejere_er_bedre
            )
            
            for placering, (chauffoer, _) in enumerate(sorterede_chauffoerer, 1):
                placeringer[chauffoer][noegletal] = placering
        
        # Beregn samlet score og sorter efter vægtkorrigeret forbrug ved uafgjort
        samlet_ranking = []
        for chauffoer, placering_dict in placeringer.items():
            samlet_score = sum(placering_dict.values())
            vaegt_forbrug = noegletal_data[chauffoer].get('Vægtkorrigeret Forbrug', float('inf'))
            samlet_ranking.append((chauffoer, samlet_score, vaegt_forbrug))
        
        # Sorter efter samlet score og derefter vægtkorrigeret forbrug
        samlet_ranking.sort(key=lambda x: (x[1], x[2]))
        
        # Opret tabel for samlet rangering
        tabel = self.doc.add_table(rows=1, cols=7)
        tabel.style = 'Table Grid'
        
        # Formater header
        header_celler = tabel.rows[0].cells
        for cell in header_celler:
            cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="E0E0E0"/>'.format(nsdecls('w'))))
        
        headers = ['Placering', 'Chauffør', 'Samlet Score', 'Tomgang', 'Fartpilot', 'Motorbremse', 'Påløbsdrift']
        for i, header in enumerate(headers):
            header_celler[i].text = header
        
        # Tilføj data
        for index, (chauffoer, score, _) in enumerate(samlet_ranking, 1):
            row_cells = tabel.add_row().cells
            row_cells[0].text = str(index)
            row_cells[1].text = chauffoer
            row_cells[2].text = str(score)
            
            # Tilføj individuelle placeringer
            row_cells[3].text = str(placeringer[chauffoer]['Tomgangsprocent'])
            row_cells[4].text = str(placeringer[chauffoer]['Fartpilot Andel'])
            row_cells[5].text = str(placeringer[chauffoer]['Motorbremse Andel'])
            row_cells[6].text = str(placeringer[chauffoer]['Påløbsdrift Andel'])
            
            # Farvemarkering for top 3
            if index <= 3:
                for cell in row_cells:
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
        
        self.doc.add_paragraph()
        self.doc.add_page_break()

    def opret_forside(self, group_name=None):
        """Opretter rapportens forside med titel og dato"""
        db_navn = os.path.basename(self.db_path)
        dele = db_navn.replace('.db', '').split('_')
        maaned = dele[2].capitalize()
        aar = dele[3]
        
        titel = self.doc.add_paragraph()
        if group_name:
            titel_tekst = titel.add_run(f'Fiskelogistik\nChaufførrapport\n{group_name}')
        else:
            titel_tekst = titel.add_run('Fiskelogistik\nChaufførrapport')
        titel_tekst.font.size = Pt(36)
        titel_tekst.font.bold = True
        titel_tekst.font.color.rgb = RGBColor(30, 144, 255)
        titel.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        periode = self.doc.add_paragraph()
        periode_tekst = periode.add_run(f'{maaned} {aar}')
        periode_tekst.font.size = Pt(24)
        periode.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        genereret_dato = self.doc.add_paragraph()
        genereret_tekst = genereret_dato.add_run(f'Genereret: {datetime.now().strftime("%d-%m-%Y %H:%M")}')
        genereret_dato.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_page_break()

    def tilfoej_sektion_overskrift(self, tekst):
        """Tilføjer en formateret sektionsoverskrift"""
        overskrift = self.doc.add_paragraph()
        overskrift_tekst = overskrift.add_run(tekst)
        overskrift_tekst.font.bold = True
        overskrift_tekst.font.size = Pt(14)
        overskrift_tekst.font.color.rgb = RGBColor(30, 144, 255)

    def opret_data_tabel(self, data, kolonner, titel):
        """Opretter en tabel med specificerede data"""
        self.tilfoej_sektion_overskrift(titel)
        
        # Tilføj selve tabellen
        tabel = self.doc.add_table(rows=1, cols=2)
        tabel.style = 'Table Grid'
        tabel.autofit = True
        
        # Tilføj headers
        header_celler = tabel.rows[0].cells
        header_celler[0].text = 'Parameter'
        header_celler[1].text = 'Værdi'
        
        # Formater headers med grå baggrund
        for cell in header_celler:
            cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="E0E0E0"/>'.format(nsdecls('w'))))
        
        # Tilføj data
        for kolonne in kolonner:
            række_celler = tabel.add_row().cells
            række_celler[0].text = kolonne
            værdi = data.get(kolonne, 'N/A')
            række_celler[1].text = str(værdi)
        
        self.doc.add_paragraph()

    def opret_noegletal_tabel(self, data, include_explanations=True):
        """Opretter en tabel med nøgletal og deres forklaringer"""
        self.tilfoej_sektion_overskrift("Nøgletal")
        
        # Definer nøgletal og deres mål
        noegletal_config = {
            'Tomgangsprocent': {
                'navn': 'Tomgang',
                'format': '{:.1f}%',
                'maal': 'Under 5%',
                'forklaring': 'Procentdel af total motordriftstid brugt i tomgang. Lavere er bedre.'
            },
            'Fartpilot Andel': {
                'navn': 'Fartpilot Anvendelse',
                'format': '{:.1f}%',
                'maal': 'Over 66,5%',
                'forklaring': 'Procentdel af køretid hvor fartpilot er aktivt. Højere er bedre.'
            },
            'Motorbremse Andel': {
                'navn': 'Brug af Motorbremse',
                'format': '{:.1f}%',
                'maal': 'Over 56%',
                'forklaring': 'Procentdel af total bremsning udført med motorbremse. Højere er bedre.'
            },
            'Påløbsdrift Andel': {
                'navn': 'Påløbsdrift',
                'format': '{:.1f}%',
                'maal': 'Over 7%',
                'forklaring': 'Procentdel af køretid i påløbsdrift. Højere er bedre.'
            },
            'Diesel Effektivitet': {
                'navn': 'Brændstofeffektivitet',
                'format': '{:.2f} km/l',
                'maal': '',
                'forklaring': 'Antal kilometer kørt per liter brændstof. Højere er bedre.'
            },
            'Vægtkorrigeret Forbrug': {
                'navn': 'Vægtkorrigeret Forbrug',
                'format': '{:.2f} l/100km/t',
                'maal': '',
                'forklaring': 'Brændstofforbrug justeret for lastens vægt. Lavere er bedre.'
            },
            'Overspeed Andel': {
                'navn': 'Hastighedsoverskridelser',
                'format': '{:.1f}%',
                'maal': '',
                'forklaring': 'Procentdel af køretid over hastighedsgrænsen. Lavere er bedre.'
            }
        }
        
        noegletal = self.beregn_noegletal(data)
        if not noegletal:
            return
        
        # Find tidligere data
        db_navn = os.path.basename(self.db_path)
        dele = db_navn.replace('.db', '').split('_')
        aktuel_maaned = dele[2]
        aktuel_aar = dele[3]
        
        tidligere_db, tidligere_data, tidligere_maaned, tidligere_aar = self.find_tidligere_database_og_data(
            aktuel_maaned, 
            aktuel_aar,
            data['Chauffør']
        )
        
        tidligere_noegletal = None
        if tidligere_data:
            tidligere_noegletal = self.beregn_noegletal(tidligere_data)
        
        tabel = self.doc.add_table(rows=1, cols=5)  # Tilføjet en ekstra kolonne til mål
        tabel.style = 'Table Grid'
        tabel.autofit = True
        
        # Tilføj headers med ny rækkefølge og bedre beskrivelser
        header_celler = tabel.rows[0].cells
        header_celler[0].text = 'Parameter'
        if tidligere_maaned and tidligere_aar:
            header_celler[1].text = f'Tidligere ({tidligere_maaned} {tidligere_aar})'
        else:
            header_celler[1].text = 'Tidligere (Ingen data)'
        header_celler[2].text = 'Nuværende'
        header_celler[3].text = 'Mål'
        header_celler[4].text = 'Udvikling siden sidst'
        
        # Formater headers med grå baggrund
        for cell in header_celler:
            cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="E0E0E0"/>'.format(nsdecls('w'))))
        
        # Tilføj nøgletal til tabellen med ny rækkefølge
        for noegletal_navn, config in noegletal_config.items():
            række_celler = tabel.add_row().cells
            række_celler[0].text = config['forklaring']
            
            # Tidligere værdi
            if tidligere_noegletal:
                tidligere_værdi = tidligere_noegletal.get(noegletal_navn, 0)
                række_celler[1].text = config['format'].format(tidligere_værdi)
            else:
                række_celler[1].text = "Ingen data"
            
            # Nuværende værdi
            værdi = noegletal.get(noegletal_navn, 0)
            række_celler[2].text = config['format'].format(værdi)
            
            # Mål værdi
            række_celler[3].text = config['maal']
            
            # Forskel/udvikling
            if tidligere_noegletal:
                tidligere_værdi = tidligere_noegletal.get(noegletal_navn, 0)
                if tidligere_værdi != 0:
                    forskel_pct = ((værdi - tidligere_værdi) / tidligere_værdi) * 100
                    forskel_tekst = f"{forskel_pct:+.1f}%"
                    
                    # Bestem farve baseret på om højere er bedre
                    if noegletal_navn in ['Tomgangsprocent', 'Vægtkorrigeret Forbrug', 'Overspeed Andel']:
                        er_forbedring = forskel_pct < 0  # Lavere er bedre
                    else:
                        er_forbedring = forskel_pct > 0  # Højere er bedre
                    
                    farve = RGBColor(0, 128, 0) if er_forbedring else RGBColor(255, 0, 0)
                    
                    # Tilføj farvet tekst
                    forskel_paragraph = række_celler[4].paragraphs[0]
                    forskel_run = forskel_paragraph.add_run(forskel_tekst)
                    forskel_run.font.color.rgb = farve
                else:
                    række_celler[4].text = "N/A"
            else:
                række_celler[4].text = "N/A"
        
        self.doc.add_paragraph()
        
        # Kun tilføj forklaringer hvis det er specificeret
        if include_explanations:
            forklaring = self.doc.add_paragraph()
            forklaring.add_run(
                "Nøgletallene giver et overblik over de vigtigste præstationsindikatorer:\n\n"
                "• Tomgang: Andel af tiden hvor motoren kører uden at køretøjet bevæger sig. "
                "En lavere procent er bedre, da tomgang bruger unødvendigt brændstof. Mål: Under 5%\n\n"
                "• Fartpilot Anvendelse: Hvor meget fartpiloten bruges ved hastigheder over 50 km/t. "
                "En højere procent er bedre, da det giver mere jævn og økonomisk kørsel. Mål: Over 66,5%\n\n"
                "• Brug af Motorbremse: Hvor meget motorbremsning bruges i forhold til normale bremser. "
                "En højere procent er bedre, da det reducerer slid på bremserne og kan genindvinde energi. Mål: Over 56%\n\n"
                "• Påløbsdrift: Hvor meget køretøjet ruller uden motorens trækkraft. "
                "En højere procent er bedre, da det sparer brændstof. Mål: Over 7%\n\n"
                "• Diesel Effektivitet: Antal kilometer kørt per liter diesel. "
                "En højere værdi er bedre, da det betyder lavere brændstofforbrug.\n\n"
                "• Vægtkorrigeret Forbrug: Brændstofforbrug justeret efter køretøjets vægt. "
                "Giver mulighed for fair sammenligning mellem forskellige læs.\n\n"
                "• Overspeed Andel: Hvor meget der køres over hastighedsgrænsen. "
                "En lavere procent er bedre af hensyn til sikkerhed og brændstofforbrug.\n"
            ).font.size = Pt(11)

    def opret_performance_rangering(self, kvalificerede_chauffoerer):
        """Opretter performancerangering for hver nøgletalskategori"""
        self.tilfoej_sektion_overskrift("Performance Rangering")
        
        intro_tekst = self.doc.add_paragraph()
        intro_run = intro_tekst.add_run(
            "Nedenstående tabeller viser rangeringen af chauffører baseret på forskellige "
            "performancemålinger. Rangeringen tager højde for om højere eller lavere værdier "
            "er optimale for hvert parameter. Grøn markering indikerer at målet er opfyldt."
        )
        intro_run.font.size = Pt(11)
        
        # Hent data for alle kvalificerede chauffører
        chauffoer_data = {}
        noegletal_data = {}
        
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        for chauffoer, _ in kvalificerede_chauffoerer:
            cursor.execute('SELECT * FROM chauffør_data_data WHERE Chauffør = ?', (chauffoer,))
            data = dict(zip([col[0] for col in cursor.description], cursor.fetchone()))
            chauffoer_data[chauffoer] = data
            noegletal_data[chauffoer] = self.beregn_noegletal(data)
        
        conn.close()
        
        # Definer nøgletal og deres optimeringsmål
        noegletal_optimering = {
            'Tomgangsprocent': ('Tomgang', False, '%', 'Lavere er bedre - Mål: Under 5% - Indikerer effektiv udnyttelse af køretøjet', 5),
            'Fartpilot Andel': ('Fartpilot Anvendelse', True, '%', 'Højere er bedre - Mål: Over 66,5% - Bidrager til jævn og økonomisk kørsel', 66.5),
            'Motorbremse Andel': ('Brug af Motorbremse', True, '%', 'Højere er bedre - Mål: Over 56% - Sparer på bremserne og reducerer brændstofforbrug', 56),
            'Påløbsdrift Andel': ('Påløbsdrift', True, '%', 'Højere er bedre - Mål: Over 7% - Indikerer effektiv udnyttelse af motorbremsning', 7),
            'Diesel Effektivitet': ('Brændstofeffektivitet', True, 'km/l', 'Højere er bedre - Indikerer effektivt brændstofforbrug', None),
            'Vægtkorrigeret Forbrug': ('Vægtkorrigeret Forbrug', False, 'l/100km/t', 'Lavere er bedre - Indikerer effektivt forbrug i forhold til vægt', None),
            'Overspeed Andel': ('Overspeed', False, '%', 'Lavere er bedre - Indikerer overholdelse af hastighedsgrænser', None)
        }
        
        # Opret en tabel for hvert nøgletal på en ny side
        for noegletal, (titel, hoejere_er_bedre, enhed, beskrivelse, maal) in noegletal_optimering.items():
            # Tilføj titel og beskrivelse
            self.doc.add_paragraph().add_run(f"\n{titel}").bold = True
            beskrivelse_para = self.doc.add_paragraph()
            beskrivelse_run = beskrivelse_para.add_run(beskrivelse)
            beskrivelse_run.font.size = Pt(10)
            beskrivelse_run.font.italic = True
            
            # Sorter chauffører efter nøgletal
            sorterede_chauffoerer = sorted(
                noegletal_data.items(),
                key=lambda x: x[1].get(noegletal, 0),
                reverse=hoejere_er_bedre
            )
            
            # Opret tabel
            tabel = self.doc.add_table(rows=1, cols=3)
            tabel.style = 'Table Grid'
            
            # Formater header
            header_celler = tabel.rows[0].cells
            for cell in header_celler:
                cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="E0E0E0"/>'.format(nsdecls('w'))))
            
            header_celler[0].text = 'Placering'
            header_celler[1].text = 'Chauffør'
            header_celler[2].text = f'Score ({enhed})'
            
            # Tilføj data med farvemarkering
            for index, (chauffoer, data) in enumerate(sorterede_chauffoerer, 1):
                row_cells = tabel.add_row().cells
                row_cells[0].text = str(index)
                row_cells[1].text = chauffoer
                
                score = data.get(noegletal, 0)
                score_tekst = f"{score:.1f}{enhed}"
                
                score_para = row_cells[2].paragraphs[0]
                score_run = score_para.add_run(score_tekst)
                
                # Marker med grøn hvis målet er opfyldt
                if maal is not None:
                    if (hoejere_er_bedre and score >= maal) or (not hoejere_er_bedre and score <= maal):
                        score_run.font.color.rgb = RGBColor(0, 128, 0)  # Grøn for opfyldt mål
            
            self.doc.add_paragraph()
            self.doc.add_page_break()  # Tilføj sideskift efter hver rangeringstabel

    def find_tidligere_database_og_data(self, aktuel_maaned, aktuel_aar, chauffoer, max_maaneder_tilbage=12):
        """Finder den seneste tidligere database hvor chaufføren findes"""
        try:
            # Dansk til nummer konvertering for måneder
            maaneder = {
                'januar': 1, 'februar': 2, 'marts': 3, 'april': 4, 
                'maj': 5, 'juni': 6, 'juli': 7, 'august': 8, 
                'september': 9, 'oktober': 10, 'november': 11, 'december': 12
            }
            
            # Nummer til dansk konvertering
            maaneder_reverse = {v: k for k, v in maaneder.items()}
            
            # Konverter aktuel måned til nummer
            aktuel_maaned_num = maaneder[aktuel_maaned.lower()]
            aktuel_dato = datetime(int(aktuel_aar), aktuel_maaned_num, 1)
            
            # Gennemgå de sidste 12 måneder
            for i in range(1, max_maaneder_tilbage + 1):
                tidligere_dato = aktuel_dato - relativedelta(months=i)
                tidligere_maaned = maaneder_reverse[tidligere_dato.month]
                tidligere_aar = str(tidligere_dato.year)
                
                # Tjek alle potentielle database navne
                database_navne = [
                    f"chauffør_data_{tidligere_maaned}_{tidligere_aar}.db",
                    f"chaufførdata_{tidligere_maaned}_{tidligere_aar}.db",
                    f"chauffør_{tidligere_maaned}_{tidligere_aar}.db"
                ]
                
                for db_navn in database_navne:
                    db_sti = os.path.join('databases', db_navn)
                    if os.path.exists(db_sti):
                        try:
                            conn = sqlite3.connect(db_sti)
                            cursor = conn.cursor()
                            
                            # Prøv forskellige tabelnavne
                            tabel_navne = ['chauffør_data_data', 'chaufførdata', 'chauffør_data']
                            
                            for tabel in tabel_navne:
                                try:
                                    cursor.execute(f'SELECT * FROM {tabel} WHERE Chauffør = ?', (chauffoer,))
                                    tidligere_data = cursor.fetchone()
                                    if tidligere_data:
                                        tidligere_data_dict = dict(zip([col[0] for col in cursor.description], tidligere_data))
                                        conn.close()
                                        return db_sti, tidligere_data_dict, tidligere_maaned.capitalize(), tidligere_aar
                                except sqlite3.OperationalError:
                                    continue
                                    
                            conn.close()
                        except Exception as e:
                            print(f"Fejl ved læsning af database {db_navn}: {str(e)}")
                            continue
                            
            return None, None, None, None
            
        except Exception as e:
            print(f"Fejl ved søgning efter tidligere database: {str(e)}")
            return None, None, None, None

    def generer_rapport(self):
        """Hovedfunktion til generering af rapporten"""
        try:
            # Opret forbindelse til databasen
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            # Opret forside
            self.opret_forside()
            
            # Find kvalificerede chauffører
            cursor.execute(f'''
                SELECT DISTINCT Chauffør, "Kørestrækning [km]"
                FROM chauffør_data_data 
                WHERE "Kørestrækning [km]" >= ?
            ''', (self.min_km,))
            
            kvalificerede_chauffoerer = cursor.fetchall()
            
            # Tilføj samlet rangering efter forsiden
            self.opret_samlet_rangering(kvalificerede_chauffoerer)
            
            # Tilføj performance rangering
            self.opret_performance_rangering(kvalificerede_chauffoerer)
            
            # Tilføj data for hver chauffør
            for chauffoer, distance in kvalificerede_chauffoerer:
                # Hent chaufførens data
                cursor.execute('''
                    SELECT * FROM chauffør_data_data 
                    WHERE Chauffør = ?
                ''', (chauffoer,))
                
                chauffoer_data = dict(zip([col[0] for col in cursor.description], cursor.fetchone()))
                
                # Tilføj chaufførnavn som overskrift
                chauffoer_overskrift = self.doc.add_heading(chauffoer, level=1)
                chauffoer_overskrift.runs[0].font.color.rgb = RGBColor(30, 144, 255)
                
                # Opret tabeller for hver datasektion
                self.opret_data_tabel(chauffoer_data, self.driftsdata_kolonner, "Driftsdata")
                self.opret_data_tabel(chauffoer_data, self.korselsdata_kolonner, "Kørselsdata")
                self.opret_data_tabel(chauffoer_data, self.tomgangsdata_kolonner, "Tomgangsdata")
                
                # Tilføj nøgletal
                self.opret_noegletal_tabel(chauffoer_data)
                
                # Tilføj sideskift mellem chauffører
                self.doc.add_page_break()
            
            # Luk databaseforbindelsen
            conn.close()
            
            # Tilføj forklaringssektion én gang efter alle data er tilføjet
            self.tilfoej_sektion_overskrift("Forklaring af Data")
            self.tilfoej_forklaringer()
            
            # Generer filnavn og gem
            db_navn = os.path.basename(self.db_path)
            dele = db_navn.replace('.db', '').split('_')
            maaned = dele[2].capitalize()
            aar = dele[3]
            tidsstempel = datetime.now().strftime("%Y%m%d_%H%M%S")
            filnavn = f"Fiskelogistik_Chaufforrapport_{maaned}_{aar}_{tidsstempel}.docx"
            
            if not os.path.exists('rapporter'):
                os.makedirs('rapporter')
            
            fuld_sti = os.path.join('rapporter', filnavn)
            self.doc.save(fuld_sti)
            
            return filnavn
            
        except Exception as e:
            raise Exception(f"Fejl ved generering af rapport: {str(e)}")

    def get_group_members(self, group_name):
        """Henter medlemmer af en specifik gruppe"""
        try:
            with DatabaseConnection('databases/settings.db') as conn:
                cursor = conn.cursor()
                cursor.execute('''
                    SELECT driver_name 
                    FROM group_members 
                    JOIN groups ON groups.id = group_members.group_id 
                    WHERE groups.name = ?
                ''', (group_name,))
                return [row[0] for row in cursor.fetchall()]
        except Exception as e:
            logging.error(f"Fejl ved hentning af gruppe medlemmer: {str(e)}")
            return []

    def generer_gruppe_rapport(self, group_name):
        """Genererer rapport for en specifik gruppe"""
        logging.info(f"Starter generering af gruppe rapport for: {group_name}")
        try:
            # Hent gruppe medlemmer
            with DatabaseConnection('databases/settings.db') as conn:
                cursor = conn.cursor()
                cursor.execute('''
                    SELECT driver_name 
                    FROM group_members 
                    JOIN groups ON groups.id = group_members.group_id 
                    WHERE groups.name = ?
                ''', (group_name,))
                group_members = [row[0] for row in cursor.fetchall()]
                logging.info(f"Fandt {len(group_members)} medlemmer i gruppen")
            
            if not group_members:
                logging.warning(f"Ingen medlemmer fundet i gruppen: {group_name}")
                raise Exception("Ingen medlemmer fundet i gruppen")
            
            # Log hver handling i processen
            logging.info("Henter kvalificerede chauffører fra databasen")
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            cursor.execute(f'''
                SELECT DISTINCT Chauffør 
                FROM chauffør_data_data 
                WHERE "Kørestrækning [km]" >= ?
            ''', (self.min_km,))
            
            qualified_drivers = [row[0] for row in cursor.fetchall()]
            logging.info(f"Fandt {len(qualified_drivers)} kvalificerede chauffører")
            
            # Find kvalificerede gruppe medlemmer
            valid_members = [driver for driver in group_members 
                            if driver in qualified_drivers]
            
            if not valid_members:
                raise Exception("Ingen kvalificerede chauffører i gruppen for denne periode")
            
            # Opret rapport med kun gruppe medlemmer
            self.doc = Document()
            
            # Opret forside med gruppe information
            self.opret_forside(group_name=group_name)
            
            # Tilføj samlet rangering kun for gruppe medlemmer
            self.opret_samlet_rangering(
                [(driver, None) for driver in valid_members]
            )
            
            # Tilføj performance rangering
            self.opret_performance_rangering(
                [(driver, None) for driver in valid_members]
            )
            
            # Tilføj individuelle chauffør sider
            for chauffoer in valid_members:
                cursor.execute('''
                    SELECT * FROM chauffør_data_data 
                    WHERE Chauffør = ?
                ''', (chauffoer,))
                
                chauffoer_data = dict(zip([col[0] for col in cursor.description], 
                                        cursor.fetchone()))
                
                # Tilføj chaufførnavn som overskrift
                chauffoer_overskrift = self.doc.add_heading(chauffoer, level=1)
                chauffoer_overskrift.runs[0].font.color.rgb = RGBColor(30, 144, 255)
                
                # Opret tabeller for hver datasektion
                self.opret_data_tabel(chauffoer_data, self.driftsdata_kolonner, 
                                    "Driftsdata")
                self.opret_data_tabel(chauffoer_data, self.korselsdata_kolonner, 
                                    "Kørselsdata")
                self.opret_data_tabel(chauffoer_data, self.tomgangsdata_kolonner, 
                                    "Tomgangsdata")
                
                # Tilføj nøgletal
                self.opret_noegletal_tabel(chauffoer_data)
                
                # Tilføj sideskift mellem chauffører
                self.doc.add_page_break()
            
            conn.close()
            
            # Generer filnavn med gruppe navn
            db_navn = os.path.basename(self.db_path)
            dele = db_navn.replace('.db', '').split('_')
            maaned = dele[2].capitalize()
            aar = dele[3]
            tidsstempel = datetime.now().strftime("%Y%m%d_%H%M%S")
            
            filnavn = f"Fiskelogistik_Gruppe_{group_name}_{maaned}_{aar}_{tidsstempel}.docx"
            
            # Gem dokumentet
            if not os.path.exists('rapporter'):
                os.makedirs('rapporter')
            
            fuld_sti = os.path.join('rapporter', filnavn)
            self.doc.save(fuld_sti)
            
            return filnavn
            
        except Exception as e:
            logging.error(f"Fejl ved generering af gruppe rapport: {str(e)}")
            raise

    def generer_individuel_rapport(self, chauffoer_navn):
        """Genererer rapport for en specifik chauffør"""
        try:
            # Hent alle kvalificerede chauffører
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            cursor.execute(f'''
                SELECT DISTINCT Chauffør, "Kørestrækning [km]"
                FROM chauffør_data_data 
                WHERE "Kørestrækning [km]" >= ?
            ''', (self.min_km,))
            
            kvalificerede_chauffoerer = cursor.fetchall()
            
            # Tjek om den valgte chauffør er kvalificeret
            if chauffoer_navn not in [c[0] for c in kvalificerede_chauffoerer]:
                raise Exception(f"Ingen kvalificeret data fundet for {chauffoer_navn} i denne periode")
            
            # Hent chaufførens data
            cursor.execute('''
                SELECT * FROM chauffør_data_data 
                WHERE Chauffør = ?
            ''', (chauffoer_navn,))
            
            chauffoer_data = dict(zip([col[0] for col in cursor.description], cursor.fetchone()))
            
            # Opret ny rapport
            self.doc = Document()
            
            # Opret forside med chauffør navn
            self.opret_forside(chauffoer_navn)
            
            # Tilføj samlet rangering med alle kvalificerede chauffører
            self.opret_samlet_rangering(kvalificerede_chauffoerer)
            
            # Tilføj performance rangering med alle kvalificerede chauffører
            self.opret_performance_rangering(kvalificerede_chauffoerer)
            
            # Tilføj den individuelle chaufførs data
            chauffoer_overskrift = self.doc.add_heading(chauffoer_navn, level=1)
            chauffoer_overskrift.runs[0].font.color.rgb = RGBColor(30, 144, 255)
            
            # Opret tabeller for hver datasektion
            self.opret_data_tabel(chauffoer_data, self.driftsdata_kolonner, "Driftsdata")
            self.opret_data_tabel(chauffoer_data, self.korselsdata_kolonner, "Kørselsdata")
            self.opret_data_tabel(chauffoer_data, self.tomgangsdata_kolonner, "Tomgangsdata")
            
            # Tilføj nøgletal
            self.opret_noegletal_tabel(chauffoer_data)
            
            # Tilføj forklaringssektion til sidst
            self.doc.add_page_break()
            self.tilfoej_sektion_overskrift("Forklaring af Data")
            self.tilfoej_forklaringer()
            
            # Generer filnavn
            db_navn = os.path.basename(self.db_path)
            dele = db_navn.replace('.db', '').split('_')
            maaned = dele[2].capitalize()
            aar = dele[3]
            tidsstempel = datetime.now().strftime("%Y%m%d_%H%M%S")
            
            # Fjern ugyldige filnavn karakterer fra chaufførnavn
            sikkert_navn = "".join(c for c in chauffoer_navn if c.isalnum() or c in (' ', '-', '_'))
            
            filnavn = f"Fiskelogistik_Chauffør_{sikkert_navn}_{maaned}_{aar}_{tidsstempel}.docx"
            
            # Gem dokumentet
            if not os.path.exists('rapporter'):
                os.makedirs('rapporter')
            
            fuld_sti = os.path.join('rapporter', filnavn)
            self.doc.save(fuld_sti)
            
            conn.close()
            return filnavn
            
        except Exception as e:
            raise Exception(f"Fejl ved generering af individuel rapport: {str(e)}")

    def generer_individuelle_rapporter(self):
        """Genererer individuelle rapporter for alle kvalificerede chauffører"""
        try:
            # Hent alle kvalificerede chauffører
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            cursor.execute(f'''
                SELECT DISTINCT Chauffør, "Kørestrækning [km]"
                FROM chauffør_data_data 
                WHERE "Kørestrækning [km]" >= ?
            ''', (self.min_km,))
            
            kvalificerede_chauffoerer = cursor.fetchall()
            
            if not kvalificerede_chauffoerer:
                raise Exception("Ingen kvalificerede chauffører fundet")
            
            generated_filenames = []
            
            # Generer rapport for hver kvalificeret chauffør
            for chauffoer, _ in kvalificerede_chauffoerer:
                # Hent chaufførens data
                cursor.execute('''
                    SELECT * FROM chauffør_data_data 
                    WHERE Chauffør = ?
                ''', (chauffoer,))
                
                chauffoer_data = dict(zip([col[0] for col in cursor.description], cursor.fetchone()))
                
                # Opret ny rapport
                self.doc = Document()
                
                # Opret forside med chauffør navn
                self.opret_forside(chauffoer)
                
                # Tilføj samlet rangering med alle kvalificerede chauffører
                self.opret_samlet_rangering(kvalificerede_chauffoerer)
                
                # Tilføj performance rangering med alle kvalificerede chauffører
                self.opret_performance_rangering(kvalificerede_chauffoerer)
                
                # Tilføj den individuelle chaufførs data
                chauffoer_overskrift = self.doc.add_heading(chauffoer, level=1)
                chauffoer_overskrift.runs[0].font.color.rgb = RGBColor(30, 144, 255)
                
                # Opret tabeller for hver datasektion
                self.opret_data_tabel(chauffoer_data, self.driftsdata_kolonner, "Driftsdata")
                self.opret_data_tabel(chauffoer_data, self.korselsdata_kolonner, "Kørselsdata")
                self.opret_data_tabel(chauffoer_data, self.tomgangsdata_kolonner, "Tomgangsdata")
                
                # Tilføj nøgletal
                self.opret_noegletal_tabel(chauffoer_data)
                
                # Tilføj forklaringssektion til sidst
                self.doc.add_page_break()
                self.tilfoej_sektion_overskrift("Forklaring af Data")
                self.tilfoej_forklaringer()
                
                # Generer filnavn
                db_navn = os.path.basename(self.db_path)
                dele = db_navn.replace('.db', '').split('_')
                maaned = dele[2].capitalize()
                aar = dele[3]
                tidsstempel = datetime.now().strftime("%Y%m%d_%H%M%S")
                
                # Fjern ugyldige filnavn karakterer fra chaufførnavn
                sikkert_navn = "".join(c for c in chauffoer if c.isalnum() or c in (' ', '-', '_'))
                
                filnavn = f"Fiskelogistik_Chauffør_{sikkert_navn}_{maaned}_{aar}_{tidsstempel}.docx"
                
                # Gem dokumentet
                if not os.path.exists('rapporter'):
                    os.makedirs('rapporter')
                
                fuld_sti = os.path.join('rapporter', filnavn)
                self.doc.save(fuld_sti)
                
                generated_filenames.append(filnavn)
            
            conn.close()
            return generated_filenames
            
        except Exception as e:
            raise Exception(f"Fejl ved generering af individuelle rapporter: {str(e)}")

    def tilfoej_forklaringer(self):
        """Tilføjer forklaringer til rapporten"""
        # Tilføj kun én forklaring under sektionen "Forklaring af Data"
        self.tilfoej_sektion_overskrift("Forklaring af Data")
        forklaring = self.doc.add_paragraph()
        forklaring.add_run(
            "Nøgletallene giver et overblik over de vigtigste præstationsindikatorer:\n\n"
            "• Tomgang: Andel af tiden hvor motoren kører uden at køretøjet bevæger sig. En lavere procent er bedre, da tomgang bruger unødvendigt brændstof.\n\n"
            "• Fartpilot Anvendelse: Hvor meget fartpiloten bruges ved hastigheder over 50 km/t. En højere procent er bedre, da det giver mere jævn og økonomisk kørsel.\n\n"
            "• Brug af Motorbremse: Hvor meget motorbremsning bruges i forhold til normale bremser. En højere procent er bedre, da det reducerer slid på bremserne og kan genindvinde energi.\n\n"
            "• Påløbsdrift: Hvor meget køretøjet ruller uden motorens trækkraft. En højere procent er bedre, da det sparer brændstof.\n\n"
            "• Diesel Effektivitet: Antal kilometer kørt per liter diesel. En højere værdi er bedre, da det betyder lavere brændstofforbrug.\n\n"
            "• Vægtkorrigeret Forbrug: Brændstofforbrug justeret efter køretøjets vægt. Giver mulighed for fair sammenligning mellem forskellige læs.\n\n"
            "• Overspeed Andel: Hvor meget der køres over hastighedsgrænsen. En lavere procent er bedre af hensyn til sikkerhed og brændstofforbrug.\n"
        ).font.size = Pt(11)
        
    if __name__ == "__main__":
        # Test kode
        generator = WordReportGenerator("databases/chauffør_data_marts_2024.db")
        generator.generer_rapport()