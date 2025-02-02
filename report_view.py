import customtkinter as ctk
import sqlite3
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches
import pandas as pd
from fpdf import FPDF
from tkinter import messagebox
import matplotlib.pyplot as plt
from PIL import Image
import logging
from word_report import WordReportGenerator
from report_mail_window import ReportMailWindow
from database_connection import DatabaseConnection

class ReportWindow:
    def __init__(self):
        # Tilføj DPI-konfiguration som i andre vinduer
        ctk.set_widget_scaling(1.0)
        ctk.deactivate_automatic_dpi_awareness()
        
        self.root = ctk.CTk()
        self.root.title("RIO Rapport Generator")
        
        # Udskyd maksimering til efter UI er bygget
        self.root.after(500, self._finalize_window_init)  # Øget forsinkelse
        
        # Eksplicit farvedefinition
        self.colors = {
            "background": "#F5F7FA",  # Tving hex-kode istedet for farvenavn
            "primary": "#007BFF",
            "secondary": "#6C757D",
            "text_primary": "#343A40",
            "text_secondary": "#6C757D",
            "card": "#FFFFFF",
            "success": "#28A745"
        }
        
        # Variables
        self.selected_type = None
        self.selected_format = None
        self.available_databases = []
        self.selected_database = None
        
        self.setup_ui()
        
    def _finalize_window_init(self):
        """Håndterer vinduesinitialisering efter UI-load"""
        self.root.update_idletasks()  # Tvinger fuld UI-opdatering
        self.root.state("zoomed")
        logging.info("Rapportvindue fuldt initialiseret")
        
    def setup_ui(self):
        # Hovedcontainer med scrolling
        self.main_container = ctk.CTkScrollableFrame(
            self.root,
            fg_color=self.colors["background"],
            height=800  # Sæt en passende højde
        )
        self.main_container.pack(expand=True, fill="both", padx=0, pady=0)
        
        # Titel sektion
        self.create_title_section(self.main_container)
        
        # Rapport type sektion
        self.create_type_section(self.main_container)
        
        # Database vælger sektion (skjult ved start)
        self.db_frame = ctk.CTkFrame(self.main_container, fg_color=self.colors["card"])
        
        # Format vælger sektion (skjult ved start)
        self.format_frame = ctk.CTkFrame(self.main_container, fg_color=self.colors["card"])
        
        # Status label
        self.status_label = ctk.CTkLabel(
            self.main_container,
            text="",
            font=("Segoe UI", 12),
            text_color=self.colors["text_secondary"]
        )
        self.status_label.pack(pady=20)

    def create_title_section(self, parent):
        title_frame = ctk.CTkFrame(parent, fg_color="transparent")
        title_frame.pack(fill="x", pady=(20, 20), padx=40)
        
        title = ctk.CTkLabel(
            title_frame,
            text="Rapport Generator",
            font=("Segoe UI", 24, "bold"),
            text_color=self.colors["primary"]
        )
        title.pack()
        
        subtitle = ctk.CTkLabel(
            title_frame,
            text="Generer detaljerede rapporter over chauffører og køretøjer",
            font=("Segoe UI", 14),
            text_color=self.colors["text_secondary"]
        )
        subtitle.pack(pady=(5, 0))

    def create_type_section(self, parent):
        type_frame = ctk.CTkFrame(parent, fg_color="transparent")
        type_frame.pack(fill="x", padx=40, pady=10)
        
        # Container for de tre bokse (nu med gruppe rapport)
        boxes_frame = ctk.CTkFrame(type_frame, fg_color="transparent")
        boxes_frame.pack(expand=True, fill="both")
        boxes_frame.grid_columnconfigure(0, weight=1)
        boxes_frame.grid_columnconfigure(1, weight=1)
        boxes_frame.grid_columnconfigure(2, weight=1)
        
        # Samlet rapport boks
        self.create_type_box(
            boxes_frame, 
            "Samlet Rapport", 
            "Generer samlet rapport over\nalle kvalificerede chauffører",
            "Inkluderer alle chauffører\nmed minimum kørestrækning",
            0
        )
        
        # Gruppe rapport boks
        self.create_type_box(
            boxes_frame, 
            "Gruppe Rapport", 
            "Generer rapport for\nspecifik chauffør gruppe",
            "Vælg en gruppe og generer\nrapport for dens medlemmer",
            1
        )
        
        # Individuel rapport boks
        self.create_type_box(
            boxes_frame, 
            "Individuel Rapport", 
            "Generer rapport for\nenkelt chauffør",
            "Vælg en specifik chauffør\nog generer rapport",
            2
        )

    def create_type_box(self, parent, title, description, details, column):
        box_container = ctk.CTkFrame(
            parent,
            fg_color=self.colors["card"],
            corner_radius=15
        )
        box_container.grid(row=0, column=column, padx=10, pady=10, sticky="nsew")
        
        title_label = ctk.CTkLabel(
            box_container,
            text=title,
            font=("Segoe UI", 18, "bold"),
            text_color=self.colors["primary"]
        )
        title_label.pack(pady=(20, 5))
        
        desc_label = ctk.CTkLabel(
            box_container,
            text=description,
            font=("Segoe UI", 12),
            text_color=self.colors["text_secondary"]
        )
        desc_label.pack(pady=(0, 10))
        
        details_label = ctk.CTkLabel(
            box_container,
            text=details,
            font=("Segoe UI", 11),
            text_color=self.colors["text_secondary"]
        )
        details_label.pack(pady=(0, 20))
        
        select_button = ctk.CTkButton(
            box_container,
            text="Vælg",
            font=("Segoe UI", 12),
            fg_color=self.colors["primary"],
            hover_color="#1874CD",
            command=lambda t=title: self.type_selected(t)
        )
        select_button.pack(pady=(0, 20))

    def create_database_section(self):
        """Opretter database vælger sektion"""
        # Fjern eksisterende frame hvis den findes
        if hasattr(self, 'db_frame'):
            self.db_frame.pack_forget()
            
        # Opret ny frame
        self.db_frame = ctk.CTkFrame(self.main_container, fg_color=self.colors["card"])
        self.db_frame.pack(fill="x", padx=40, pady=10)
        
        # Overskrift
        header = ctk.CTkLabel(
            self.db_frame,
            text="Vælg Database",
            font=("Segoe UI", 16, "bold"),
            text_color=self.colors["primary"]
        )
        header.pack(pady=10)
        
        # Få tilgængelige databaser
        self.available_databases = self.get_available_databases()
        
        if not self.available_databases:
            no_data_label = ctk.CTkLabel(
                self.db_frame,
                text="Ingen databaser fundet for den valgte rapport type",
                font=("Segoe UI", 12),
                text_color=self.colors["text_secondary"]
            )
            no_data_label.pack(pady=10)
            return
                
        # Database vælger container
        db_select_frame = ctk.CTkFrame(self.db_frame, fg_color="transparent")
        db_select_frame.pack(pady=10, fill="x", padx=20)
        
        # Database dropdown
        db_label = ctk.CTkLabel(
            db_select_frame,
            text="Vælg Database:",
            font=("Segoe UI", 14),
            text_color=self.colors["text_primary"]
        )
        db_label.pack(pady=(0, 5))
        
        # Formatér databasenavne til mere læsbare navne
        formatted_names = [self.format_db_name(db) for db in self.available_databases]
        
        self.db_var = ctk.StringVar(value="Vælg database")
        db_dropdown = ctk.CTkOptionMenu(
            db_select_frame,
            values=formatted_names,
            variable=self.db_var,
            command=self.database_selected,
            width=300
        )
        db_dropdown.pack(pady=5)

    def create_format_section(self):
        """Opretter format vælger sektion"""
        self.format_frame.pack(fill="x", padx=40, pady=10)
        
        # Overskrift
        header = ctk.CTkLabel(
            self.format_frame,
            text="Vælg Format og Handling",
            font=("Segoe UI", 16, "bold"),
            text_color=self.colors["primary"]
        )
        header.pack(pady=10)
        
        # Format vælger container
        format_select_frame = ctk.CTkFrame(self.format_frame, fg_color="transparent")
        format_select_frame.pack(pady=10, fill="x", padx=20)
        
        # Format knapper
        button_frame = ctk.CTkFrame(format_select_frame, fg_color="transparent")
        button_frame.pack(fill="x", pady=10)
        
        # Word knap
        word_button = ctk.CTkButton(
            button_frame,
            text="Generer Word",
            command=lambda: self.format_selected("word"),
            fg_color=self.colors["primary"],
            hover_color="#1874CD"
        )
        word_button.pack(side="left", padx=5)
        
        # PDF knap
        pdf_button = ctk.CTkButton(
            button_frame,
            text="Generer PDF",
            command=lambda: self.format_selected("pdf"),
            fg_color=self.colors["primary"],
            hover_color="#1874CD"
        )
        pdf_button.pack(side="left", padx=5)
        
        # Send Mail knap
        mail_button = ctk.CTkButton(
            button_frame,
            text="Send Rapporter",
            command=self.show_mail_window,
            fg_color=self.colors["success"],
            hover_color="#1e7e34"
        )
        mail_button.pack(side="right", padx=5)

    def format_db_name(self, db_name):
        # Fjern .db extension og split på underscore
        parts = db_name.replace('.db', '').split('_')
        if len(parts) >= 4:  # chauffør_data_januar_2024.db
            data_type = parts[0].capitalize()
            month = parts[2].capitalize()
            year = parts[3]
            return f"{data_type} - {month} {year}"
        return db_name

    def database_selected(self, formatted_name):
        """Håndterer valg af database"""
        # Find den oprindelige database fil ud fra det formaterede navn
        for db in self.available_databases:
            if self.format_db_name(db) == formatted_name:
                try:
                    # Gem den valgte database sti
                    self.selected_database = os.path.join('databases', db)
                    
                    # Fjern tidligere action frame hvis den findes
                    if hasattr(self, 'action_frame'):
                        self.action_frame.pack_forget()
                    
                    # Vis handling sektion
                    self.create_action_section()
                    
                    self.status_label.configure(
                        text=f"Database valgt: {formatted_name}",
                        text_color=self.colors["text_primary"]
                    )
                    
                except Exception as e:
                    self.status_label.configure(
                        text=f"Fejl ved valg af database: {str(e)}",
                        text_color="red"
                    )
                break

    def type_selected(self, type_name):
        """Håndterer valg af rapport type"""
        self.selected_type = type_name.lower().split()[0]  # "samlet", "gruppe" eller "individuel"
        
        # Fjern tidligere frames hvis de eksisterer
        if hasattr(self, 'db_frame'):
            self.db_frame.pack_forget()
        if hasattr(self, 'action_frame'):
            self.action_frame.pack_forget()
        if hasattr(self, 'format_frame'):
            self.format_frame.pack_forget()
        if hasattr(self, 'group_frame'):
            self.group_frame.pack_forget()
        
        # Vis gruppe vælger KUN hvis gruppe rapport er valgt
        if self.selected_type == "gruppe":
            self.create_group_selector()
        
        # Opret database vælger
        self.create_database_section()
        
        self.status_label.configure(
            text=f"Valgt: {type_name}",
            text_color=self.colors["text_primary"]
        )

    def get_available_databases(self):
        """Henter tilgængelige databaser baseret på valgt rapport type"""
        if not os.path.exists('databases'):
            return []
            
        databases = []
        prefix = "chauffør_data"  # Standard prefix for chauffør databaser
        
        for file in os.listdir('databases'):
            if file.startswith(prefix) and file.endswith('.db'):
                databases.append(file)
                
        return sorted(databases, reverse=True)  # Nyeste først

    def format_selected(self, format_ext):
        self.selected_format = format_ext
        self.status_label.configure(
            text=f"Format valgt: {format_ext.upper()}",
            text_color=self.colors["text_primary"]
        )
        
        # Generer rapport når format er valgt
        self.generate_report()

    def generate_report(self):
        if self.selected_database is None:
            self.status_label.configure(
                text="Ingen database valgt",
                text_color="red"
            )
            return
            
        try:
            if not os.path.exists('rapporter'):
                os.makedirs('rapporter')
            
            # Opret WordReportGenerator
            word_generator = WordReportGenerator(self.selected_database)
            
            if self.selected_type == "samlet":
                # Generer samlet rapport
                generated_filename = word_generator.generer_rapport()
                success_message = "Samlet rapport genereret"
            elif self.selected_type == "gruppe":
                if not hasattr(self, 'selected_group') or not self.selected_group.get():
                    messagebox.showerror("Fejl", "Vælg venligst en gruppe")
                    return
                # Generer gruppe rapport
                generated_filename = word_generator.generer_gruppe_rapport(
                    self.selected_group.get()
                )
                success_message = "Gruppe rapport genereret"
            elif self.selected_type == "individuel":
                # Generer individuelle rapporter for alle kvalificerede chauffører
                generated_filenames = word_generator.generer_individuelle_rapporter()
                success_message = f"{len(generated_filenames)} individuelle rapporter genereret"
                
                if not generated_filenames:
                    messagebox.showerror("Fejl", "Ingen kvalificerede chauffører fundet")
                    return
                    
                # Vis bekræftelse med antal genererede rapporter
                messagebox.showinfo(
                    "Rapporter Genereret",
                    f"{len(generated_filenames)} individuelle rapporter er blevet gemt i mappen 'rapporter'"
                )
                return
            
            self.status_label.configure(
                text=success_message,
                text_color="green"
            )
            
            if self.selected_type != "individuel":
                messagebox.showinfo(
                    "Rapport Genereret",
                    f"Rapporten er blevet gemt som:\n{generated_filename}\n\ni mappen 'rapporter'"
                )
            
        except Exception as e:
            self.status_label.configure(
                text=f"Fejl under generering af rapport: {str(e)}",
                text_color="red"
            )
            messagebox.showerror("Fejl", f"Kunne ikke generere rapport: {str(e)}")

    def generate_word_report(self, filename, df):
        doc = Document()
        
        # Tilføj titel
        doc.add_heading(f'{self.selected_type.capitalize()} Rapport', 0)
        doc.add_paragraph(f'Genereret: {datetime.now().strftime("%d-%m-%Y %H:%M")}')
        
        # Tilføj tabel
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'
        
        # Tilføj headers
        for i, column in enumerate(df.columns):
            table.cell(0, i).text = column
            
        # Tilføj data
            for _, row in df.iterrows():
                cells = table.add_row().cells
                for i, value in enumerate(row):
                    cells[i].text = str(value)
                    
            # Tilføj grundlggende statistik
            doc.add_heading('Statistik', level=1)
            for column in df.select_dtypes(include=['float64', 'int64']).columns:
                doc.add_paragraph(
                    f'{column}:\n'
                    f'Gennemsnit: {df[column].mean():.2f}\n'
                    f'Minimum: {df[column].min():.2f}\n'
                    f'Maximum: {df[column].max():.2f}'
                )
                    
            doc.save(os.path.join('rapporter', f'{filename}.docx'))

    def generate_pdf_report(self, filename, df):
        pdf = FPDF()
        pdf.add_page()
        
        # Konfigurer font
        pdf.add_font('DejaVu', '', 'DejaVuSansCondensed.ttf', uni=True)
        pdf.set_font('DejaVu', '', 12)
        
        # Tilføj titel
        pdf.set_font('DejaVu', '', 24)
        pdf.cell(0, 10, f'{self.selected_type.capitalize()} Rapport', ln=True, align='C')
        
        # Tilføj dato
        pdf.set_font('DejaVu', '', 12)
        pdf.cell(0, 10, f'Genereret: {datetime.now().strftime("%d-%m-%Y %H:%M")}', ln=True)
        
        # Tilføj tabel
        pdf.add_page()
        col_width = pdf.w / len(df.columns)
        row_height = pdf.font_size * 1.5
        
        # Headers
        for column in df.columns:
            pdf.cell(col_width, row_height, str(column), 1)
        pdf.ln()
        
        # Data
        for _, row in df.iterrows():
            for value in row:
                pdf.cell(col_width, row_height, str(value), 1)
            pdf.ln()
            
        # Tilføj statistik
        pdf.add_page()
        pdf.set_font('DejaVu', '', 16)
        pdf.cell(0, 10, 'Statistik', ln=True)
        pdf.set_font('DejaVu', '', 12)
        
        for column in df.select_dtypes(include=['float64', 'int64']).columns:
            pdf.multi_cell(0, 10, 
                f'{column}:\n'
                f'Gennemsnit: {df[column].mean():.2f}\n'
                f'Minimum: {df[column].min():.2f}\n'
                f'Maximum: {df[column].max():.2f}\n'
            )
            
        pdf.output(os.path.join('rapporter', f'{filename}.pdf'))

    def generate_excel_report(self, filename, df):
        # Opret Excel writer
        excel_path = os.path.join('rapporter', f'{filename}.xlsx')
        writer = pd.ExcelWriter(excel_path, engine='xlsxwriter')
        
        # Gem hoveddata
        df.to_excel(writer, sheet_name='Data', index=False)
        
        # Opret statistik ark
        stats_df = pd.DataFrame()
        for column in df.select_dtypes(include=['float64', 'int64']).columns:
            stats_df[column] = [
                df[column].mean(),
                df[column].min(),
                df[column].max(),
                df[column].std()
            ]
        
        stats_df.index = ['Gennemsnit', 'Minimum', 'Maximum', 'Standardafvigelse']
        stats_df.to_excel(writer, sheet_name='Statistik')
        
        # Gem og luk
        writer.close()

    def create_group_selector(self):
        """Opretter gruppe vælger sektion"""
        # Opret gruppe frame
        self.group_frame = ctk.CTkFrame(self.main_container, fg_color=self.colors["card"])
        self.group_frame.pack(fill="x", padx=40, pady=10)
        
        # Overskrift
        header = ctk.CTkLabel(
            self.group_frame,
            text="Vælg Gruppe",
            font=("Segoe UI", 16, "bold"),
            text_color=self.colors["primary"]
        )
        header.pack(pady=10)
        
        # Gruppe vælger container
        group_select_frame = ctk.CTkFrame(self.group_frame, fg_color="transparent")
        group_select_frame.pack(pady=10, fill="x", padx=20)
        
        # Gruppe dropdown
        group_label = ctk.CTkLabel(
            group_select_frame,
            text="Vælg Gruppe:",
            font=("Segoe UI", 14),
            text_color=self.colors["text_primary"]
        )
        group_label.pack(pady=(0, 5))
        
        # Hent tilgængelige grupper
        groups = self.get_available_groups()
        group_names = [group[1] for group in groups] if groups else ["Ingen grupper fundet"]
        
        self.group_var = ctk.StringVar(value="Vælg gruppe")
        group_dropdown = ctk.CTkOptionMenu(
            group_select_frame,
            values=group_names,
            variable=self.group_var,
            width=300
        )
        group_dropdown.pack(pady=5)

    def get_available_groups(self):
        """Henter tilgængelige grupper fra databasen"""
        try:
            with DatabaseConnection('databases/settings.db') as conn:
                cursor = conn.cursor()
                cursor.execute('SELECT id, name FROM groups')
                return cursor.fetchall()
        except Exception as e:
            logging.error(f"Fejl ved hentning af grupper: {str(e)}")
            return []

    def create_action_section(self):
        """Opretter handling sektion med rapport muligheder"""
        # Fjern eksisterende frame hvis den findes
        if hasattr(self, 'action_frame'):
            self.action_frame.pack_forget()
            
        # Opret ny frame
        self.action_frame = ctk.CTkFrame(self.main_container, fg_color=self.colors["card"])
        self.action_frame.pack(fill="x", padx=40, pady=10)
        
        # Indre container med padding
        inner_frame = ctk.CTkFrame(self.action_frame, fg_color=self.colors["card"])
        inner_frame.pack(fill="x", padx=20, pady=20)
        
        # Overskrift
        header = ctk.CTkLabel(
            inner_frame,
            text="Vælg Handling",
            font=("Segoe UI", 18, "bold"),
            text_color=self.colors["primary"]
        )
        header.pack(pady=(0,15))
        
        # Beskrivelse
        description_text = "Vælg hvordan du vil håndtere rapporten."
        if self.selected_type == "individuel":
            description_text += " Du kan generere en Word/PDF fil eller sende direkte via mail."
        else:
            description_text += " Du kan generere rapporten som Word eller PDF fil."
        
        description = ctk.CTkLabel(
            inner_frame,
            text=description_text,
            font=("Segoe UI", 12),
            text_color=self.colors["text_secondary"],
            wraplength=600
        )
        description.pack(pady=(0,20))
        
        # Knap container med centrering
        button_container = ctk.CTkFrame(inner_frame, fg_color="transparent")
        button_container.pack(expand=True, fill="x")
        
        # Center frame til knapper
        button_frame = ctk.CTkFrame(button_container, fg_color="transparent")
        button_frame.pack(expand=True)
        
        # Word rapport knap
        word_button = ctk.CTkButton(
            button_frame,
            text="Generer Word Rapport",
            command=lambda: self.format_selected("word"),
            fg_color=self.colors["primary"],
            hover_color="#1874CD",
            width=200,
            height=40,
            corner_radius=8
        )
        word_button.pack(side="left", padx=10)
        
        # PDF rapport knap
        pdf_button = ctk.CTkButton(
            button_frame,
            text="Generer PDF Rapport",
            command=lambda: self.format_selected("pdf"),
            fg_color=self.colors["primary"],
            hover_color="#1874CD",
            width=200,
            height=40,
            corner_radius=8
        )
        pdf_button.pack(side="left", padx=10)
        
        # Send rapport knap - kun ved individuel rapport
        if self.selected_type == "individuel":
            send_button = ctk.CTkButton(
                button_frame,
                text="Send Rapport",
                command=self.show_mail_window,
                fg_color=self.colors["success"],
                hover_color="#1e7e34",
                width=200,
                height=40,
                corner_radius=8
            )
            send_button.pack(side="left", padx=10)

    def run(self):
        """Starter applikationen"""
        try:
            self.root.state("zoomed")
            self.root.mainloop()
        except Exception as e:
            messagebox.showerror("Fejl", f"Kunne ikke starte rapportvindue: {str(e)}")

    def destroy(self):
        """Lukker vinduet og frigør ressourcer"""
        try:
            # Destroy alle child windows først
            for widget in self.root.winfo_children():
                if isinstance(widget, ctk.CTkToplevel):
                    widget.destroy()
            
            # Destroy hovedvinduet
            self.root.destroy()
            
            # Afbryd mainloop
            self.root.quit()
            
        except Exception as e:
            print(f"Fejl ved lukning af rapport vindue: {str(e)}")

    def show_mail_window(self):
        """Viser mail sending vinduet"""
        try:
            if not self.selected_database:
                messagebox.showwarning(
                    "Advarsel",
                    "Vælg venligst en database først"
                )
                return
                
            # Tjek om mail er konfigureret
            db = DatabaseConnection('settings.db')
            mail_config = db.get_mail_config()
            
            if not mail_config:
                response = messagebox.askyesno(
                    "Mail Ikke Konfigureret",
                    "Mail indstillinger er ikke konfigureret.\n\n" +
                    "Du skal konfigurere følgende:\n" +
                    "- SMTP server og port\n" +
                    "- Email og adgangskode\n" +
                    "- Test email adresse\n\n" +
                    "Vil du konfigurere mail indstillingerne nu?"
                )
                
                if response:
                    # Importer og åbn settings vinduet
                    from settings_view import SettingsWindow
                    settings = SettingsWindow()
                    settings.tabview.set("Mail")  # Skift til mail fanen
                    return
                return
                
            # Opret og vis mail vinduet
            from report_mail_window import ReportMailWindow
            mail_window = ReportMailWindow(
                self.root,
                self.selected_database,
                self.selected_type,
                self.selected_group if hasattr(self, 'selected_group') else None,
                self.selected_driver if hasattr(self, 'selected_driver') else None
            )
            mail_window.run()
            
        except Exception as e:
            error_msg = str(e)
            if "no such column: smtp_server" in error_msg:
                error_msg = "Mail systemet er ikke korrekt konfigureret. Gå til Indstillinger -> Mail for at konfigurere det."
            
            logging.error(f"Fejl ved åbning af mail vindue: {error_msg}")
            messagebox.showerror(
                "Fejl",
                f"Kunne ikke åbne mail vindue: {str(e)}"
            )

if __name__ == "__main__":
    app = ReportWindow()
    app.run()